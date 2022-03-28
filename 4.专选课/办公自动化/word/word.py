# -*- coding:utf-8 -*-
'''
@File    :   word.py
@Time    :   2022/03/28 08:57:37
@Author  :   hyh
@Version :   1.0
@Contact :   1360895771@qq.com
@Desc    :   word课上练习
'''
# here put the import lib
import docx
import openpyxl

doc = docx.Document('libai.docx')
doc.add_heading('静夜思',0)
doc.add_heading('李白',1)
doc.save("静夜思.docx")
